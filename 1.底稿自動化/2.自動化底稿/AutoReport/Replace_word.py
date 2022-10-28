import xlwings as xl
# import os
# import win32com.client
from openpyxl import load_workbook
import docx
import warnings
warnings.filterwarnings("ignore")




def CallVBA():
    app = xl.App(visible=True)
    wb = app.books.open('工作底稿.xlsm')
    wb.macro('main')()
    app.quit()











def info_update(doc,old_info, new_info):
    '''此函數用於批量替換合同中需要替換的信息
    doc:文件
    old_info和new_info：原文字和需要替換的新文字
    '''
    
    

    #讀取段落中的所有run，找到需替換的信息進行替換
    for para in doc.paragraphs: #

        for run in para.runs:
            run.text = run.text.replace(old_info, new_info) #替換信息


    #讀取表格中的所有單元格，找到需替換的信息進行替換
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info) #替換信息



def Replace_Word():
    doc = docx.Document('sample.docx')
    wb = load_workbook('工作底稿.xlsm')
    sheet  = wb['查核稽核報告']
    Range = ['B','C','D']
    
    for columns in Range:
        for i in range(50,  2,-1):
            try:
                info_update(doc, old_info=f'{columns}{str(i)}', new_info=sheet[f'{columns}{str(i)}'].value)
            except:
                pass
    
    doc.save('查核稽核報告.docx')
    print(f'執行完成')





