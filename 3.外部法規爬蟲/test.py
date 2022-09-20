import glob
import aspose.words as aw
import docx
import pandas as pd
import os
import re

path =r"\\192.168.248.10\ia\主管機關相關法規\外籍移工匯兌公司防制洗錢及打擊資恐注意事項範本(含指引及態樣)"
files = glob.glob(path+'\*.pdf')

files=files[2]


tital = files.split('\\')[-1].replace('.pdf','')
convert = pd.read_csv(os.path.join('datasource','conver_1.csv')).sort_index(ascending=False)
old_word = convert['old'].tolist()
new_word = [str(_) for _ in convert['new'].tolist()]

# 後續改loop

# load the PDF file
doc = aw.Document(files)

# convert PDF to Word DOCX format
filename = "TMP.docx"
doc.save(filename)

doc = docx.Document(filename)

text = ''.join([_.text for _ in doc.paragraphs])


for old,new in zip(old_word,new_word):
    text = text.replace(old, new)

text = text.replace('第','@@').replace('條 ','&&')
law= [_.replace('@','第').replace('&','條') for _ in  re.findall('@\d*&',text)]
content_etl=re.split('@\d*&',text)
content_etl = [_.replace('@','第').replace('&','_') for _ in content_etl if len(_)>2]

df = pd.DataFrame({
    '條號' : law,
    '內容' : content_etl
    })


df['法規名稱'] = tital
df = df[['法規名稱','條號','內容']]


#%%
import os
import glob
import PyPDF2
import pandas as pd
import re





path =r"\\192.168.248.10\ia\主管機關相關法規\金融機構整合電子化支付端末設備注意事項"
files = glob.glob(path+'\*.pdf')

for file in files:
    tital = file.split('\\')[-1].replace('.pdf','')
    convert = pd.read_csv(os.path.join('datasource','conver_1.csv')).sort_index(ascending=False)
    # creating an object 
    file = open(file, 'rb')
    
    # creating a pdf reader object
    fileReader = PyPDF2.PdfFileReader(file)
    
    
    content = []
    for page in fileReader.pages:
        # pprint.pprint(page.extractText())
        content.append(page.extractText())
    
    
    content_location="\n".join(content).replace('第一條', '@第一條').replace(' ', '').split('@')
    content = content_location[1]
    
    # content=content.replace('\n', '!!')
    
    old_word = convert['old'].tolist()
    new_word = [str(_) for _ in convert['new'].tolist()]
    # new_word = convert['new'].tolist()
    
    for old,new in zip(old_word,new_word):
        content = content.replace(old, new)
        # print(old,new)
    
    
    content = content.replace('第','@@').replace('條\n\n','&&')
    law= [_.replace('@','第').replace('&','條') for _ in  re.findall('@\d*&',content)]
    content_etl=re.split('@\d*&',content)
    content_etl = [_.replace('@','第').replace('&','_') for _ in content_etl if len(_)>2]
    
    df = pd.DataFrame({
        '條號' : law,
        '內容' : content_etl
        })
    
    
    df['法規名稱'] = tital
    df = df[['法規名稱','條號','內容']]


